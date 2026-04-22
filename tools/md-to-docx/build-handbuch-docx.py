#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
build-handbuch-docx.py

Konvertiert docs/Benutzerhandbuch.md in eine .docx-Datei mit eingebetteten
Screenshots. Die Quelle ist ASCII-transliteriert (ae/oe/ue/ss); dieses Skript
uebersetzt die betroffenen Woerter anhand einer expliziten Wort-Whitelist in
echte deutsche Umlaute und Eszett (ä, ö, ü, Ä, Ö, Ü, ß).

Aufruf:
    python tools/md-to-docx/build-handbuch-docx.py

Ausgabe:
    docs/Benutzerhandbuch.docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# tools/md-to-docx/build-handbuch-docx.py → parents[2] == Repo-Root
REPO_ROOT = Path(__file__).resolve().parents[2]
MD_PATH = REPO_ROOT / "docs" / "Benutzerhandbuch.md"
DOCX_PATH = REPO_ROOT / "docs" / "Benutzerhandbuch.docx"
IMG_BASE = REPO_ROOT / "docs"

MAX_IMG_WIDTH_CM = 15.5


# =============================================================================
# Rueck-Transliteration: explizite Wort-Mapping-Tabelle.
# Jedes Schluessel-Wort taucht als Token im Handbuch auf; Nichtvorkommen in
# dieser Tabelle bedeutet: das Wort bleibt unveraendert (z. B. echte
# u+e-Folgen wie "neue", "aktuell", "manuell", "Grauer", "Blauer" oder
# englische Fachtokens wie "Session", "password", "Permissions").
# =============================================================================
WORD_MAP: dict[str, str] = {
    # A
    "abgeschlossen": "abgeschlossen",
    "Abgeschlossen": "Abgeschlossen",
    "Abgeschlossene": "Abgeschlossene",
    "Abhaengig": "Abhängig",
    "abschliessen": "abschließen",
    "abschliesst": "abschließt",
    "Abschluss": "Abschluss",
    "Abschlussfest": "Abschlussfest",
    "Adresse": "Adresse",
    "Adressen": "Adressen",
    "aelteste": "älteste",
    "aenderbar": "änderbar",
    "Aendern": "Ändern",
    "aendern": "ändern",
    "aendert": "ändert",
    "Aenderung": "Änderung",
    "Aenderungen": "Änderungen",
    "Aenderungsprotokoll": "Änderungsprotokoll",
    "Aktivitaet": "Aktivität",
    "aktuell": "aktuell",
    "aktuelle": "aktuelle",
    "aktuellen": "aktuellen",
    "Aktueller": "Aktueller",
    "angepasst": "angepasst",
    "anpassen": "anpassen",
    "anschliessender": "anschließender",
    "anstossen": "anstoßen",
    "Antraege": "Anträge",
    "Arbeitsstunden": "Arbeitsstunden",
    "arbeitsstunden": "arbeitsstunden",
    "Aufschluesselung": "Aufschlüsselung",
    "ausgefuellt": "ausgefüllt",
    "ausgewaehlt": "ausgewählt",
    "Ausschalten": "Ausschalten",
    "ausschliesslich": "ausschließlich",
    "ausserdem": "außerdem",
    "auswaehlbar": "auswählbar",
    # B
    "Bearbeitungssperre": "Bearbeitungssperre",
    "Bearbeitungssperren": "Bearbeitungssperren",
    "Begruendung": "Begründung",
    "Begruessungsnachricht": "Begrüßungsnachricht",
    "benoetigen": "benötigen",
    "Benoetigt": "Benötigt",
    "benoetigten": "benötigten",
    "benoetigter": "benötigter",
    "Benutzermenue": "Benutzermenü",
    "Berechtigungssystem": "Berechtigungssystem",
    "bestaetigen": "bestätigen",
    "Bestaetigen": "Bestätigen",
    "bestaetigt": "bestätigt",
    "Bestaetigungscode": "Bestätigungscode",
    "Blauer": "Blauer",
    # D
    "dass": "dass",
    "Datenueberschreibung": "Datenüberschreibung",
    "Dauer": "Dauer",
    "desselben": "desselben",
    "dessen": "dessen",
    "duerfen": "dürfen",
    "durchfuehren": "durchführen",
    "durchgefuehrt": "durchgeführt",
    "durchlaeuft": "durchläuft",
    # E
    "Eigentuemer": "Eigentümer",
    "einfuegen": "einfügen",
    "Einrichtungsschluessel": "Einrichtungsschlüssel",
    "Einrichtungsseite": "Einrichtungsseite",
    "Einsaetze": "Einsätze",
    "Einsaetzen": "Einsätzen",
    "einschliesslich": "einschließlich",
    "Einschraenkung": "Einschränkung",
    "Eintraege": "Einträge",
    "eintraege": "einträge",
    "Eintraegen": "Einträgen",
    "Eintragseigentuemers": "Eintragseigentümers",
    "Eintragsstatus": "Eintragsstatus",
    "enthaelt": "enthält",
    "Entwurfsstatus": "Entwurfsstatus",
    "erfassen": "erfassen",
    "erfasser": "erfasser",
    "Erfasser": "Erfasser",
    "Erfassung": "Erfassung",
    "Erfassungsformular": "Erfassungsformular",
    "Erfassungssystem": "Erfassungssystem",
    "erfuellen": "erfüllen",
    "Erfuellt": "Erfüllt",
    "erfuellte": "erfüllte",
    "Erfuellung": "Erfüllung",
    "Erfuellungsgrad": "Erfüllungsgrad",
    "erhaelt": "erhält",
    "erhoehte": "erhöhte",
    "ermoeglicht": "ermöglicht",
    # F
    "faehig": "fähig",
    "Faellen": "Fällen",
    "fliessen": "fließen",
    "Fuegen": "Fügen",
    "fuehrt": "führt",
    "Fuellen": "Füllen",
    "fuer": "für",
    "Fuer": "Für",
    # G
    "geaendert": "geändert",
    "gegenueber": "gegenüber",
    "Geloescht": "Gelöscht",
    "geloescht": "gelöscht",
    "Geloeschte": "Gelöschte",
    "geoeffnet": "geöffnet",
    "Geraet": "Gerät",
    "Geraete": "Geräte",
    "geraeuschlos": "geräuschlos",
    "geschuetzt": "geschützt",
    "Getraenke": "Getränke",
    "gewaehlte": "gewählte",
    "gewaehlten": "gewählten",
    "gewuenschte": "gewünschte",
    "gewuenschten": "gewünschten",
    "gezaehlt": "gezählt",
    "Grauer": "Grauer",
    "Gross": "Groß",
    "Grossbuchstabe": "Großbuchstabe",
    "Gruen": "Grün",
    "gruen": "grün",
    "gruenen": "grünen",
    "gruenes": "grünes",
    "Grundriss": "Grundriss",
    "gueltig": "gültig",
    "Gueltigkeit": "Gültigkeit",
    # H
    "Haekchen": "Häkchen",
    "haengen": "hängen",
    "haeufige": "häufige",
    "Haeufige": "Häufige",
    "hinzufuegen": "hinzufügen",
    "Hinzufuegen": "Hinzufügen",
    # I
    "Inaktivitaet": "Inaktivität",
    "Individuelle": "Individuelle",
    # J
    "Jaehrliches": "Jährliches",
    "jaehrliches": "jährliches",
    # K
    "klaerung": "klärung",
    "Klaerung": "Klärung",
    "Klaerungsbedarf": "Klärungsbedarf",
    "koennen": "können",
    "kuenftig": "künftig",
    "kuerzer": "kürzer",
    # L
    "laedt": "lädt",
    "laesst": "lässt",
    "laeuft": "läuft",
    "lassen": "lassen",
    "linksbuendig": "linksbündig",
    "Loeschen": "Löschen",
    "loeschen": "löschen",
    "Loescht": "Löscht",
    "Loeschung": "Löschung",
    "lueckenlos": "lückenlos",
    # M
    "Manipulationsschutz": "Manipulationsschutz",
    "Manuell": "Manuell",
    "manuell": "manuell",
    "Manuelle": "Manuelle",
    "manuelle": "manuelle",
    "manuellen": "manuellen",
    "Manuelles": "Manuelles",
    "Menue": "Menü",
    "Menuepunkte": "Menüpunkte",
    "missbraeuchliche": "missbräuchliche",
    "Missbrauch": "Missbrauch",
    "Moechten": "Möchten",
    "moechten": "möchten",
    "moeglich": "möglich",
    "Moegliche": "Mögliche",
    "Monatsuebersicht": "Monatsübersicht",
    "muessen": "müssen",
    "Muss": "Muss",
    "muss": "muss",
    # N
    "nachtraeglich": "nachträglich",
    "naechsten": "nächsten",
    "naechster": "nächster",
    "Neue": "Neue",
    "neue": "neue",
    "neuen": "neuen",
    "Neuen": "Neuen",
    "neuer": "neuer",
    "Neuer": "Neuer",
    "neueren": "neueren",
    "neues": "neues",
    "Neues": "Neues",
    "noetig": "nötig",
    # O
    "oeffentlichen": "öffentlichen",
    "Oeffnen": "Öffnen",
    "oeffnen": "öffnen",
    "oeffnet": "öffnet",
    "Oeffnet": "Öffnet",
    # P
    "passiert": "passiert",
    "passt": "passt",
    "password": "password",
    "Passwort": "Passwort",
    "passwort": "passwort",
    "Passwortaenderung": "Passwortänderung",
    "Passworts": "Passworts",
    "Permissions": "Permissions",
    "Persoenliche": "Persönliche",
    "persoenlichen": "persönlichen",
    "Pessimistische": "Pessimistische",
    "Plaetze": "Plätze",
    "Plaetzen": "Plätzen",
    "Pruefen": "Prüfen",
    "pruefen": "prüfen",
    "pruefer": "prüfer",
    "Pruefer": "Prüfer",
    "prueferliste": "prüferliste",
    "pruefliste": "prüfliste",
    "Pruefliste": "Prüfliste",
    "prueft": "prüft",
    "Pruefung": "Prüfung",
    "Pruefungsinformationen": "Prüfungsinformationen",
    "Pruefungszwecke": "Prüfungszwecke",
    # R
    "rechtsbuendig": "rechtsbündig",
    "Rollenuebersicht": "Rollenübersicht",
    "rollenuebersicht": "rollenübersicht",
    "rueckfrage": "rückfrage",
    "Rueckfrage": "Rückfrage",
    "Rueckfragen": "Rückfragen",
    "rueckfragen": "rückfragen",
    "Ruecknahme": "Rücknahme",
    # S
    "Schaltflaeche": "Schaltfläche",
    "schliessen": "schließen",
    "Schluessel": "Schlüssel",
    "schwarzweissen": "schwarzweißen",
    "selbststaendig": "selbstständig",
    "Session": "Session",
    "Sicherheitsmassnahme": "Sicherheitsmaßnahme",
    "Sicherheitsstufe": "Sicherheitsstufe",
    "sodass": "sodass",
    "spaeter": "später",
    "Sperrdauer": "Sperrdauer",
    "standardmaessig": "standardmäßig",
    "Standardmaessig": "Standardmäßig",
    "strasse": "straße",
    "Strasse": "Straße",
    "Stundenerfassung": "Stundenerfassung",
    "Stundenerfassungsformular": "Stundenerfassungsformular",
    # T
    "Taetigkeit": "Tätigkeit",
    # U
    "Ueber": "Über",
    "ueber": "über",
    "Ueberblick": "Überblick",
    "Uebergaenge": "Übergänge",
    "uebergebene": "übergebene",
    "uebergeordnete": "übergeordnete",
    "uebernehmen": "übernehmen",
    "Uebernommen": "Übernommen",
    "uebernommen": "übernommen",
    "Ueberschreiben": "Überschreiben",
    "ueberschreiben": "überschreiben",
    "ueberschreibt": "überschreibt",
    "ueberschrieben": "überschrieben",
    "Uebersicht": "Übersicht",
    "uebersicht": "übersicht",
    "uebersichtlich": "übersichtlich",
    "Uebersichtsliste": "Übersichtsliste",
    "uebersichtsliste": "übersichtsliste",
    "Uebersichtsseite": "Übersichtsseite",
    "Uebersprungen": "Übersprungen",
    "umfassende": "umfassende",
    "unabhaengig": "unabhängig",
    "unberuehrt": "unberührt",
    "ungeprueften": "ungeprüften",
    "ungueltig": "ungültig",
    "Ungueltige": "Ungültige",
    "Ungueltiger": "Ungültiger",
    "Unterstuetzt": "Unterstützt",
    "unterstuetzt": "unterstützt",
    "Unterstuetzte": "Unterstützte",
    "urspruengliche": "ursprüngliche",
    # V
    "veraendert": "verändert",
    "Verbesserte": "Verbesserte",
    "verfuegbar": "verfügbar",
    "Verfuegung": "Verfügung",
    "vergessen": "vergessen",
    "Vergessen": "Vergessen",
    "Verschluesselung": "Verschlüsselung",
    "vollstaendig": "vollständig",
    "Vollstaendige": "Vollständige",
    "vollstaendigen": "vollständigen",
    "Vollstaendigen": "Vollständigen",
    "Vollstaendiger": "Vollständiger",
    "vorausgewaehlt": "vorausgewählt",
    "Voraussetzung": "Voraussetzung",
    "Vorschlaege": "Vorschläge",
    "voruebergehend": "vorübergehend",
    # W
    "waehlen": "wählen",
    "Waehlen": "Wählen",
    "Waende": "Wände",
    # Z
    "zaehlen": "zählen",
    "zuerst": "zuerst",
    "zugelassen": "zugelassen",
    "zukuenftige": "zukünftige",
    "zunaechst": "zunächst",
    "zurueck": "zurück",
    "zurueckblaettern": "zurückblättern",
    "zurueckgezogen": "zurückgezogen",
    "zurueckkehrt": "zurückkehrt",
    "zuruecknehmen": "zurücknehmen",
    "zuruecksenden": "zurücksenden",
    "zuruecksetzen": "zurücksetzen",
    "Zuruecksetzen": "Zurücksetzen",
    "Zurueckziehen": "Zurückziehen",
    "zurueckziehen": "zurückziehen",
    "zurueckzusetzen": "zurückzusetzen",
    "Zusaetzlich": "Zusätzlich",
    "zusaetzlich": "zusätzlich",
    "zusaetzliche": "zusätzliche",
    "Zusaetzliche": "Zusätzliche",
    "zusammenfassung": "zusammenfassung",
    "Zusammenfassung": "Zusammenfassung",
    "Zusammenfassungskarten": "Zusammenfassungskarten",
    "Zuverlaessiger": "Zuverlässiger",
    "Zuverlaessigkeit": "Zuverlässigkeit",
}


_WORD_RE = re.compile(r"[A-Za-z][A-Za-z]+")


def retransliterate(text: str) -> str:
    """Wort-weise Ersetzung. Woerter ohne Eintrag in WORD_MAP bleiben gleich."""
    def repl(match: re.Match[str]) -> str:
        w = match.group(0)
        return WORD_MAP.get(w, w)
    return _WORD_RE.sub(repl, text)


# =============================================================================
# Markdown-Regex
# =============================================================================
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*?)\s*$")
IMG_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")
LIST_RE = re.compile(r"^(\s*)[-*+]\s+(.*)$")
OLIST_RE = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
HR_RE = re.compile(r"^\s*(?:-{3,}|\*{3,}|_{3,})\s*$")
QUOTE_RE = re.compile(r"^>\s?(.*)$")
TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")


def add_horizontal_line(paragraph) -> None:
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_inline_runs(paragraph, text: str) -> None:
    """Zerlegt Text in (bold/italic/code)-Segmente. Inline-Code wird NICHT
    rueck-transliteriert, alles andere ja."""
    tokens: list[tuple[str, set[str]]] = [(text, set())]

    def apply(regex: re.Pattern[str], flag: str) -> None:
        new_tokens: list[tuple[str, set[str]]] = []
        for txt, flags in tokens:
            if flag in flags:
                new_tokens.append((txt, flags))
                continue
            last = 0
            for m in regex.finditer(txt):
                if m.start() > last:
                    new_tokens.append((txt[last:m.start()], set(flags)))
                new_tokens.append((m.group(1), flags | {flag}))
                last = m.end()
            if last < len(txt):
                new_tokens.append((txt[last:], set(flags)))
        tokens[:] = new_tokens

    apply(INLINE_CODE_RE, "code")
    apply(BOLD_RE, "bold")
    apply(ITALIC_RE, "italic")

    for txt, flags in tokens:
        if not txt:
            continue
        display = txt if "code" in flags else retransliterate(txt)
        run = paragraph.add_run(display)
        if "code" in flags:
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xA0, 0x26, 0x26)
        if "bold" in flags:
            run.bold = True
        if "italic" in flags:
            run.italic = True


def embed_image(doc: Document, rel_path: str, alt: str) -> None:
    img_path = (IMG_BASE / rel_path).resolve()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not img_path.is_file():
        run = p.add_run(f"[Bild fehlt: {rel_path}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
        return
    try:
        p.add_run().add_picture(str(img_path), width=Cm(MAX_IMG_WIDTH_CM))
    except Exception as exc:  # pragma: no cover
        run = p.add_run(f"[Bildfehler {rel_path}: {exc}]")
        run.italic = True
    if alt:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(retransliterate(alt))
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def parse_table_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    cols = max(len(header), *(len(r) for r in rows)) if rows else len(header)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    for idx in range(cols):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(retransliterate(header[idx]) if idx < len(header) else "")
        run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx in range(cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            if c_idx < len(row):
                add_inline_runs(para, row[c_idx])
    doc.add_paragraph()


def set_document_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for level in range(1, 5):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0x1F, 0x3B, 0x6E)


def convert() -> None:
    if not MD_PATH.is_file():
        print(f"FEHLER: {MD_PATH} nicht gefunden.", file=sys.stderr)
        sys.exit(2)

    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    set_document_defaults(doc)

    i = 0
    in_code = False
    code_buf: list[str] = []

    def flush_code() -> None:
        if not code_buf:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_buf))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)
        code_buf.clear()

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if "|" in line and i + 1 < len(lines) and TABLE_SEP_RE.match(lines[i + 1]):
            header = parse_table_row(line)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, header, rows)
            continue

        if HR_RE.match(line):
            p = doc.add_paragraph()
            add_horizontal_line(p)
            i += 1
            continue

        m_h = HEADING_RE.match(line)
        if m_h:
            level = min(len(m_h.group(1)), 4)
            p = doc.add_heading(level=level)
            add_inline_runs(p, m_h.group(2))
            i += 1
            continue

        stripped = line.strip()
        m_img = IMG_RE.fullmatch(stripped) if stripped else None
        if m_img:
            embed_image(doc, m_img.group(2), m_img.group(1))
            i += 1
            continue

        if IMG_RE.search(line):
            pos = 0
            for m in IMG_RE.finditer(line):
                pre = line[pos:m.start()].strip()
                if pre:
                    p = doc.add_paragraph()
                    add_inline_runs(p, pre)
                embed_image(doc, m.group(2), m.group(1))
                pos = m.end()
            rest = line[pos:].strip()
            if rest:
                p = doc.add_paragraph()
                add_inline_runs(p, rest)
            i += 1
            continue

        m_q = QUOTE_RE.match(line)
        if m_q:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            add_inline_runs(p, m_q.group(1))
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        m_l = LIST_RE.match(line)
        if m_l:
            indent = len(m_l.group(1)) // 2
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.6 + 0.5 * indent)
            add_inline_runs(p, m_l.group(2))
            i += 1
            continue

        m_ol = OLIST_RE.match(line)
        if m_ol:
            indent = len(m_ol.group(1)) // 2
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.6 + 0.5 * indent)
            add_inline_runs(p, m_ol.group(3))
            i += 1
            continue

        if stripped == "":
            i += 1
            continue

        block_lines = [line]
        j = i + 1
        while j < len(lines):
            nxt = lines[j]
            stripped_nxt = nxt.strip()
            if stripped_nxt == "":
                break
            if HEADING_RE.match(nxt):
                break
            if HR_RE.match(nxt):
                break
            if LIST_RE.match(nxt) or OLIST_RE.match(nxt):
                break
            if QUOTE_RE.match(nxt):
                break
            if nxt.strip().startswith("```"):
                break
            if IMG_RE.fullmatch(stripped_nxt):
                break
            if "|" in nxt and j + 1 < len(lines) and TABLE_SEP_RE.match(lines[j + 1]):
                break
            block_lines.append(nxt)
            j += 1
        paragraph_text = " ".join(l.strip() for l in block_lines)
        p = doc.add_paragraph()
        add_inline_runs(p, paragraph_text)
        i = j

    if in_code:
        flush_code()

    doc.save(DOCX_PATH)
    print(f"OK: geschrieben -> {DOCX_PATH}")
    print(f"    Groesse: {DOCX_PATH.stat().st_size // 1024} KB")


def audit_unmapped() -> None:
    """Nennt Woerter im Dokument, die Kandidatenmuster tragen aber keinen
    Whitelist-Eintrag haben. Hilft beim Pflegen der Mapping-Tabelle."""
    text = MD_PATH.read_text(encoding="utf-8")
    text = re.sub(r"```.*?```", "", text, flags=re.DOTALL)
    text = re.sub(r"`[^`]+`", "", text)
    pattern = re.compile(r"[A-Za-z][A-Za-z]+")
    seen: set[str] = set()
    for m in pattern.finditer(text):
        w = m.group(0)
        if re.search(r"(ae|oe|ue|Ae|Oe|Ue|ss|Ss)", w) and w not in WORD_MAP:
            seen.add(w)
    if seen:
        print("Unmapped candidate words:")
        for w in sorted(seen, key=str.lower):
            print(f"  {w}")


if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "--audit":
        audit_unmapped()
    else:
        convert()
